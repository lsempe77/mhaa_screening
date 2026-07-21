"""Generate a docx answer-sheet version of the scope-decisions memo for the review team.

Produces strongminds/scope_decisions_for_ZS.docx with explicit answer fields the team
can fill in (checkbox-style options + 'Selected option' line + comments box).
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = Path(__file__).resolve().parent / "scope_decisions_for_ZS.docx"

RECORDS = {
    # Decision 1 — biological correlates as determinants
    "D1_excluded": [
        ("130320840", "MDD face-processing network connectivity", "Neuroimaging — brain connectivity during facial expression processing"),
        ("130321825", "Inter-Brain Synchrony and Psychological Conditions", "Hyperscanning / inter-brain synchrony as a correlate of depression"),
        ("130326298", "Genetic markers of stress generation in depression", "5-HTTLPR / OXTR / HPA-axis genetic correlates"),
        ("130327136", "Mendelian Randomization for psychiatric diseases incl. MDD", "Genetic instrument / drug-target MR for MDD"),
    ],
    "D1_included": [
        ("130312828", "Depression in the Iranian elderly: prevalence & meta-analysis", "Prevalence + social planning in an LMIC sub-population"),
        ("130322190", "Prevalence of psychiatric disorders among tribal populations of India", "Prevalence in a marginalized LMIC sub-population"),
    ],
    # Decision 2 — sub-population scope
    "D2_excluded": [
        ("Prisoners / incarcerated adults", "130313498, 130317150", "2"),
        ("University / medical students", "130312643, 130321297", "2"),
        ("Refugees resettled in HIC", "130313819", "1"),
        ("Disease-specific cohorts (epilepsy, CKD, post-stroke)", "130353034*, 130340196", "2"),
        ("Older adults in specific care settings", "130326299, 130335744", "2"),
        ("Parents of infants / parent-infant dyads", "130335744, 130346770*", "2"),
    ],
    "D2_included": [
        ("Older adults: dementia+depression, cognitive impairment+depression", "kept", "—"),
        ("Heart failure + CBT for depression", "kept", "—"),
        ("IPV in pregnancy with depression", "kept", "—"),
        ("Post-Ebola psychosocial support", "kept", "—"),
    ],
}


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def add_answer_block(doc, decision_no, options, recommended_idx, recommendation_text):
    p = doc.add_paragraph()
    r = p.add_run("YOUR DECISION (please tick ONE option):")
    r.bold = True
    r.font.size = Pt(10)
    for i, opt in enumerate(options):
        bp = doc.add_paragraph(style="List Bullet")
        mark = "[ ]"
        run = bp.add_run(f"{mark}  Option {chr(97+i)}: {opt}")
        run.font.size = Pt(10)
        if i == recommended_idx:
            rec = bp.add_run("   (recommended)")
            rec.italic = True
            rec.font.color.rgb = RGBColor(0x44, 0x77, 0xAA)
            rec.font.size = Pt(9)
    sel = doc.add_paragraph()
    sel.add_run("Selected option: ").bold = True
    sel.add_run("______________________")
    com = doc.add_paragraph()
    com.add_run("Comments / rationale: ").bold = True
    com.add_run("\n________________________________________________________________________")
    com.add_run("\n________________________________________________________________________")
    note = doc.add_paragraph()
    nr = note.add_run(f"Recommendation if unsure: {recommendation_text}")
    nr.italic = True
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)


def main():
    doc = Document()
    # base font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    h = doc.add_heading("Protocol-scope decisions — ULCM TAS screening", level=0)
    sub = doc.add_paragraph()
    sub.add_run("For: StrongMinds ULCM review team (ZS)\nFrom: TAS-screening iteration (v1.6 \u2192 v1.7/v1.8)\nDate: 2026-07-21").italic = True

    # Instructions
    doc.add_heading("What you need to do", level=1)
    instr = doc.add_paragraph()
    instr.add_run(
        "Three scope questions are blocking the next prompt iteration (v1.8). "
        "Each question shows the concrete records that force it, a set of options, and a recommendation. "
        "Please "
    )
    r = instr.add_run("tick ONE option per question")
    r.bold = True
    instr.add_run(
        " (Option a/b/c), add a one-line rationale if helpful, and return this document. "
        "Your answers will be encoded as explicit exclusion rules in the screener prompts. "
        "No screening re-run is needed from your side \u2014 just the three decisions."
    )

    note = doc.add_paragraph()
    nr = note.add_run(
        "Context: after cleaning 10 confirmed GT label errors, the tool sits at "
        "sensitivity 0.716 / specificity 0.917 / \u03ba 0.584 against the cleaned reference. "
        "The remaining errors are split between (a) fixable prompt issues \u2014 addressed in v1.7, already running \u2014 "
        "and (b) scope rules the GT applies but the protocol does not state. This memo covers (b). "
        "Your decisions here determine whether v1.8 can approach \u03ba 0.70."
    )
    nr.italic = True
    nr.font.size = Pt(9)
    nr.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ---- Decision 1 ----
    doc.add_heading("Decision 1 \u2014 Are biological correlates \u201cdeterminants\u201d for RQ1?", level=1)
    doc.add_paragraph(
        "RQ1 covers \u201cdrivers/risk factors for adult depression.\u201d The written criterion is broad enough to include "
        "any association with depression. But the GT excludes a cluster of records that study biological correlates of "
        "depression (neuroimaging, biomarkers, genetic polymorphisms), while including social/psychological/economic "
        "risk factors. The boundary is not stated in the protocol."
    )

    doc.add_heading("Records the GT excludes (currently pipeline INCLUDES = false positives)", level=2)
    add_table(doc,
              ["record_id", "title", "what it studies"],
              RECORDS["D1_excluded"],
              col_widths=[1.1, 2.6, 2.9])

    doc.add_heading("Records the GT includes (tool correctly keeps)", level=2)
    add_table(doc,
              ["record_id", "title", "what it studies"],
              RECORDS["D1_included"],
              col_widths=[1.1, 2.6, 2.9])

    doc.add_paragraph("Decision needed: which class of records counts as RQ1 \u201cdeterminants\u201d?")
    add_answer_block(doc, 1,
        options=[
            "Only modifiable / social / psychological / economic risk factors and prevalence. Biological / neurobiological / genetic mechanism studies are OUT of scope.",
            "All correlates of depression count, including biomarkers, neuroimaging and genetics. Keep them IN scope.",
            "Biological correlates are IN scope only when they inform intervention or policy (e.g. a biomarker validated for screening), not as pure mechanism research.",
        ],
        recommended_idx=0,
        recommendation_text="Option (a) \u2014 matches the GT's implicit practice and the review's applied (Ultra-Low-Cost) focus.")

    doc.add_paragraph()

    # ---- Decision 2 ----
    doc.add_heading("Decision 2 \u2014 Which adult sub-populations are in-scope?", level=1)
    doc.add_paragraph(
        "The written population criterion says \u201cadults 18+ with depression/CMD/distress; perinatal women eligible; "
        "mixed adult/adolescent retained unless clearly adolescent-only.\u201d The GT, however, excludes a set of adult "
        "sub-populations on \u201cpopulation\u201d even though they are adults with a depression focus. The excluding "
        "principle is not written anywhere."
    )

    doc.add_heading("Sub-populations the GT excludes on population (pipeline INCLUDES = false positives)", level=2)
    add_table(doc,
              ["sub-population", "example record_ids", "count"],
              RECORDS["D2_excluded"],
              col_widths=[3.0, 2.2, 1.0])

    doc.add_paragraph(
        "* 130353034 (epilepsy+depression, Ethiopia) and 130338268 (post-stroke aphasia) were human-confirmed as INCLUDE "
        "in your review of the 12 GT-error candidates, so the GT's \u201cdisease-specific cohort\u201d exclusion is not "
        "consistent. This is exactly the fuzziness that makes a written rule necessary."
    )

    doc.add_heading("Sub-populations the GT includes (correctly kept) \u2014 the boundary is real but unstated", level=2)
    add_table(doc,
              ["sub-population / case", "status", "\u2014"],
              RECORDS["D2_included"],
              col_widths=[3.0, 1.5, 1.5])

    doc.add_paragraph("Decision needed: please confirm which adult sub-populations are OUT of scope regardless of depression focus, by answering each sub-question below.")

    # Sub-questions with per-item checkboxes
    sub_qs = [
        ("Prisoners / incarcerated populations", ["IN scope", "OUT of scope"]),
        ("University / medical students (your note on 130341241 suggested \u201clate adolescence 18-24 is includable\u201d \u2014 does this extend to all student populations, or only when the age range is clearly 18+?)", ["IN scope", "OUT of scope", "IN only if clearly 18+"]),
        ("Refugees resettled in HIC (e.g. Southeast Asian refugees in the US) \u2014 refugees in LMIC are clearly IN; the question is the resettled-in-HIC case", ["IN scope", "OUT of scope"]),
        ("Disease-specific cohorts (epilepsy, CKD, HIV, post-stroke) \u2014 IN when the intervention/study targets depression, OUT when depression is merely measured alongside the disease?", ["IN always (if adult + depression focus)", "IN only when depression is the target", "OUT always"]),
        ("Parents of infants / parent-infant dyads \u2014 IN when the parent's depression is the target, OUT when the infant/attachment is the target?", ["IN only when parent depression is target", "OUT always", "IN always"]),
    ]
    for i, (q, opts) in enumerate(sub_qs, 1):
        p = doc.add_paragraph()
        r = p.add_run(f"2.{i}  {q}")
        r.bold = True
        r.font.size = Pt(10)
        for opt in opts:
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run(f"[ ]  {opt}").font.size = Pt(10)
        sp = doc.add_paragraph()
        sp.add_run("Selected: ").bold = True
        sp.add_run("______________________")
        doc.add_paragraph()

    # Final over-arching option for decision 2
    p = doc.add_paragraph()
    r = p.add_run("Alternatively, if you prefer ONE rule covering all sub-populations, tick here instead of answering 2.1\u20132.5 individually:")
    r.bold = True
    doc.add_paragraph("[ ]  Adopt the \u201cdepression-target rule\u201d: a sub-population is IN-scope when depression is the target of the study (intervention, determinants, or measurement); OUT-of-scope when depression is merely measured alongside a different primary subject (the disease, the infant, the incarceration, the student experience).  (recommended)", style="List Bullet")

    doc.add_paragraph()

    # ---- Decision 3 ----
    doc.add_heading("Decision 3 \u2014 Does RQ18 cover generic health-status instruments?", level=1)
    doc.add_paragraph(
        "From your note on 130377408 (\u201cEQ-5D population norms\u2026 could be a fit for RQ18. Need a second opinion\u201d): "
        "the EQ-5D is a generic health-status instrument whose dimensions include an anxiety/depression item. The "
        "question is whether RQ18 \u201cvalidity/reliability of depression measurement tools\u201d covers:"
    )
    add_answer_block(doc, 3,
        options=[
            "Depression-specific instruments only (PHQ-9, EPDS, BDI, CES-D, HSCL, K10, SRQ-20, etc.). Exclude generic instruments even if they have an anxiety/depression dimension.",
            "Any instrument with a depression-relevant dimension when the study validates that dimension specifically (e.g. include EQ-5D's AD dimension if the study is about its depression-measurement properties).",
        ],
        recommended_idx=0,
        recommendation_text="Option (a) \u2014 RQ18 is about depression measurement; a generic health-status instrument is not a \u201cdepression measure\u201d even if it touches the construct.")

    doc.add_paragraph()

    # ---- What these decisions unlock ----
    doc.add_heading("What these decisions unlock", level=1)
    doc.add_paragraph("Once decided, these scope rules will be encoded as explicit FAIL signals in v1.8. Estimated effect on the cleaned-GT metrics:")
    add_table(doc,
              ["decision", "est. FP reduction", "est. FN change", "est. \u03ba \u0394"],
              [
                  ["1 (biological correlates)", "\u22124 to \u22126", "0", "+0.03"],
                  ["2 (sub-population scope)", "\u22126 to \u22128", "0 to \u22122", "+0.04"],
                  ["3 (RQ18 instrument scope)", "\u22121", "0", "+0.005"],
              ],
              col_widths=[2.6, 1.4, 1.2, 1.0])
    doc.add_paragraph(
        "Combined with the v1.7 prompt edits (router fix + Criterion 4 reframing, already running), the target after "
        "v1.8 is sensitivity ~0.86 / specificity ~0.93 / \u03ba ~0.66\u20130.68 \u2014 approaching but not crossing the 0.70 "
        "threshold, consistent with the finding that the residual gap is irreducible fuzziness and "
        "absent-from-abstract judgment calls (only movable by full-text screening on the uncertain bucket)."
    )

    # Return instructions
    doc.add_heading("How to return this", level=1)
    doc.add_paragraph(
        "Save the file with your initials appended (e.g. scope_decisions_for_ZS_done_ZS.docx) and drop it back in the "
        "shared folder, or reply with the three selected options (e.g. \u201c1a, 2 (depression-target rule), 3a\u201d). "
        "No need to re-screen anything \u2014 the three decisions are enough."
    )

    doc.save(OUT)
    print(f"Wrote {OUT}  ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
