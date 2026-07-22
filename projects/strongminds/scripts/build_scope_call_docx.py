"""Generate a docx answer-sheet for the depression-secondary-outcome scope call.

Produces strongminds/scope_call_secondary_outcome_for_ZS.docx with the 9 records where
ZS's Decision 2 (depression must be the target) conflicts with the GT (which INCLUDES
these depression-secondary studies). ZS decides per-record IN/OUT or picks a rule.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = Path(__file__).resolve().parent.parent / "docs" / "scope_call_secondary_outcome_for_ZS.docx"

# The 9 records (record_id, title, primary_subject, why_pipeline_excluded)
RECORDS = [
    ("130345244",
     "Persistent neuropsychiatric symptoms after COVID-19: a systematic review and meta-analysis.",
     "COVID-19 / post-COVID neuropsychiatry",
     "COVID-19 is the primary subject; depression is one of several neuropsychiatric symptoms measured."),
    ("130325044",
     "The association between mental health and obesity in postmenopausal women: A systematic review.",
     "Obesity in postmenopausal women",
     "Obesity is the primary subject; depression/anxiety are measured as associated outcomes."),
    ("130316032",
     "Bedtime procrastination and psychological distress in university students: a systematic review and meta-analysis of their association.",
     "Bedtime procrastination / sleep behavior",
     "Bedtime procrastination is the primary subject; psychological distress (incl. depression) is the associated outcome."),
    ("130338536",
     "An integrative review of potential enablers and barriers to accessing mental health services in Ghana.",
     "Healthcare access / service barriers",
     "Access to mental-health services is the primary subject; depression is not the focus (it is one condition among CMDs served)."),
    ("130324026",
     "Implementing civic engagement within mental health services in South East Asia: a systematic review and realist synthesis of current evidence.",
     "Civic engagement in service delivery",
     "Civic engagement / service co-design is the primary subject; depression is not the focus."),
    ("130346770",
     "Group-based parent training programmes for improving parental psychosocial health",
     "Parenting skills / parent training",
     "Parenting skills is the intervention target; parental depression is one of several measured outcomes."),
    ("130316284",
     "The Silent Crisis: Loneliness in Older Adults-A Critical Review of Impacts, Strategies and Path Forward.",
     "Loneliness in older adults",
     "Loneliness is the primary subject; depression is an associated impact, not the focus."),
    ("130323841",
     "Hormone Therapy, Mental Health, and Quality of Life Among Transgender People: A Systematic Review.",
     "Gender-affirming hormone therapy",
     "Hormone therapy is the primary subject; depression is one of several mental-health outcomes measured."),
    ("130377518",
     "Measuring disability in consumers of mental health services - psychometric properties of the WHODAS 2.0 in Ghana.",
     "Generic disability measurement (WHODAS)",
     "WHODAS (a generic disability instrument) is the primary subject; depression is not the focus."),
]


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            for p in cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    return t


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # Title
    doc.add_heading("Scope call \u2014 depression-secondary-outcome records", level=0)
    sub = doc.add_paragraph()
    sub.add_run(
        "For: StrongMinds ULCM review team (ZS)\n"
        "From: TAS-screening iteration (v1.8.1)\n"
        "Date: 2026-07-21\n"
        "Re: Your Decision 2 (\u201cdepression-target rule\u201d) vs the current ground truth on 9 records"
    ).italic = True

    # The conflict
    doc.add_heading("The conflict", level=1)
    p = doc.add_paragraph()
    p.add_run(
        "Your Decision 2 (scope memo, 2026-07-21) adopted the rule: "
    )
    r = p.add_run("\u201ca sub-population is IN-scope when depression is the target of the study; "
                  "OUT-of-scope when depression is merely measured alongside a different primary subject "
                  "(the disease, the infant, the incarceration, the student experience).\u201d")
    r.bold = True
    p.add_run(
        " The v1.8.1 screener applies this rule strictly. But the current ground truth INCLUDES 9 records "
        "where depression is plainly NOT the primary subject \u2014 the study is about COVID, obesity, "
        "parenting, loneliness, hormone therapy, healthcare access, etc., and depression is one measured "
        "outcome among several. The pipeline correctly excludes these per your rule; the GT includes them. "
        "So these 9 are false negatives against the GT, but correct excludes per your Decision 2."
    )

    p2 = doc.add_paragraph()
    p2.add_run("This is the one remaining structural disagreement between the tool and the GT. ").bold = True
    p2.add_run(
        "Your decision here determines whether (a) we fix the GT to match your rule "
        "(all 9 \u2192 EXCLUDE \u2192 tool is right, sensitivity rises ~+0.13), or "
        "(b) we relax the rule to keep depression-mentioning-but-secondary studies IN "
        "(the 9 stay INCLUDE \u2192 tool must be loosened, but FPs will rise because the loosening is hard to bound)."
    )

    doc.add_paragraph()

    # The 9 records
    doc.add_heading("The 9 records", level=1)
    doc.add_paragraph(
        "Each record below is currently GT-INCLUDE but pipeline-EXCLUDE (vote 0.0\u20130.3, i.e. the model is "
        "confident it should be excluded). For each, the \u201cprimary subject\u201d column states what the study "
        "is actually about; \u201cwhy pipeline excluded\u201d states the model\u2019s reasoning."
    )

    for rid, title, subject, why in RECORDS:
        doc.add_heading(f"{rid} \u2014 {title}", level=2)
        sp = doc.add_paragraph()
        sp.add_run("Primary subject: ").bold = True
        sp.add_run(subject)
        wp = doc.add_paragraph()
        wp.add_run("Why pipeline excluded: ").bold = True
        wp.add_run(why)
        ap = doc.add_paragraph()
        ap.add_run("Your call: ").bold = True
        ap.add_run("[ ]  KEEP INCLUDE (depression-mentioning-but-secondary studies are IN-scope)     ")
        ap.add_run("[ ]  FLIP TO EXCLUDE (your Decision 2 applies; tool is right)     ")
        ap.add_run("[ ]  UNSURE \u2014 needs full-text")
        cp = doc.add_paragraph()
        cp.add_run("Comment: ").bold = True
        cp.add_run("______________________________________________________________")
        doc.add_paragraph()

    # Over-arching rule option
    doc.add_heading("OR \u2014 pick one rule covering all 9", level=1)
    doc.add_paragraph(
        "If you prefer not to decide record-by-record, tick ONE of these rules. The rule will be applied to "
        "all 9 (and encoded into the v1.8.2 prompt for future records of this type):"
    )
    doc.add_paragraph(
        "[ ]  Option A (strict \u2014 recommended, matches your Decision 2): Apply the depression-target rule strictly. "
        "All 9 \u2192 EXCLUDE. The GT was over-inclusive on depression-mentioning-but-secondary studies; we fix the GT. "
        "Effect: sensitivity rises ~0.703 \u2192 ~0.83, \u03ba rises ~0.66 \u2192 ~0.74 (crosses the 0.70 threshold).",
        style="List Bullet")
    doc.add_paragraph(
        "[ ]  Option B (loose): Depression-mentioning-but-secondary studies are IN-scope when the primary subject is "
        "mental-health-adjacent (loneliness, parenting psychosocial health, healthcare access, disability measurement). "
        "Keep these as INCLUDE; tell the tool to stop excluding them. Effect: sensitivity rises ~0.703 \u2192 ~0.83, but "
        "FPs will rise by a similar amount (the loosening is hard to bound \u2014 the model cannot reliably distinguish "
        "\u201cmental-health-adjacent\u201d from \u201cnon-mental-health\u201d primary subjects from the abstract).",
        style="List Bullet")
    doc.add_paragraph(
        "[ ]  Option C (per-record): I have answered each record individually above.",
        style="List Bullet")

    doc.add_paragraph()

    # What each option does
    doc.add_heading("What each option means for the metrics", level=1)
    add_table(doc,
              ["option", "sensitivity", "specificity", "\u03ba", "thresholds"],
              [
                  ["A (strict)", "~0.83", "unchanged (~0.95)", "~0.74", "\u03ba crosses 0.70 \u2705; sens still < 0.95 \u274c"],
                  ["B (loose)", "~0.83", "drops (~0.88)", "~0.60", "\u03ba falls; sens still < 0.95 \u274c"],
                  ["C (per-record)", "depends", "depends", "depends", "depends on the mix"],
              ],
              col_widths=[1.4, 1.3, 1.4, 1.0, 2.0])
    doc.add_paragraph(
        "Note: sensitivity 0.95 stays unreachable under any option. The residual ~13 FNs after Option A are "
        "genuinely ambiguous boundary cases (adolescent age ranges, non-ULCM modalities the GT includes "
        "inconsistently) where the deciding fact is not in the title/abstract. Only full-text screening on the "
        "uncertain bucket moves sensitivity further \u2014 consistent with the \u00a713 oracle result (perfect "
        "tie-breaking reaches sens 0.808, \u03ba 0.727)."
    )

    # Return instructions
    doc.add_heading("How to return this", level=1)
    doc.add_paragraph(
        "Save the file with your initials appended (e.g. scope_call_secondary_outcome_for_ZS_done_ZS.docx) and "
        "drop it back in the shared folder, or reply with the selected option (A / B / C) and any per-record flips. "
        "No re-screening needed from your side \u2014 the three decisions are enough to encode and re-run."
    )

    doc.save(OUT)
    print(f"Wrote {OUT}  ({OUT.stat().st_size} bytes)")


if __name__ == "__main__":
    main()
